# src/project/doc_reader.py

from pathlib import Path
from typing import List, Optional
import logging
import pandas as pd
import docx
import pypdf
import json

from .pydantic_models import Document, DocumentType

logger = logging.getLogger(__name__)

class DocumentReader:
    """A class to read documents from various file types."""

    SUPPORTED_FORMATS = {
        ".pdf": DocumentType.PDF,
        ".txt": DocumentType.TXT,
        ".json": DocumentType.JSON,
        ".docx": DocumentType.DOCX,
        ".csv": DocumentType.CSV,
        ".tsv": DocumentType.TSV,
    }

    @staticmethod
    def find_files(input_path: Path) -> List[Path]:
        """Finds all supported files in the input directory."""
        return [p for p in input_path.glob("**/*") if p.suffix.lower() in DocumentReader.SUPPORTED_FORMATS]

    @staticmethod
    def read_file(file_path: Path) -> Optional[Document]:
        """Reads a single file and returns a Document object."""
        doc_type = DocumentReader.SUPPORTED_FORMATS.get(file_path.suffix.lower())
        if not doc_type:
            logger.warning(f"Unsupported file type: {file_path.suffix}")
            return None

        try:
            doc_id = file_path.stem
            metadata = {"source": str(file_path), "file_name": file_path.name}
            content = ""

            if doc_type == DocumentType.PDF:
                with open(file_path, "rb") as f:
                    pdf_reader = pypdf.PdfReader(f)
                    content = "\n".join([page.extract_text() for page in pdf_reader.pages])
            elif doc_type in [DocumentType.TXT, DocumentType.JSON, DocumentType.CSV, DocumentType.TSV]:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
            elif doc_type == DocumentType.DOCX:
                doc = docx.Document(file_path)
                content = "\n".join([para.text for para in doc.paragraphs])

            return Document(
                id=doc_id,
                doc_name=file_path.name,
                content=content,
                document_type=doc_type,
                metadata=metadata,
                source=str(file_path),
                title=file_path.name
            )
        except Exception as e:
            logger.error(f"Error reading {file_path}: {e}")
            return None